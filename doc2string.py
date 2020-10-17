import docx
import nltk
import re
from nltk import word_tokenize

class Document():

  document = None
  string = None 
  tokens = None 
  words  = None 
  types  = None 
  token_ratio = None

  def __init__(self, path):
    self.document = docx.Document(path)
    self.string = self.build_string()
    self.tokens = self.build_tokens()
    self.words  = self.remove_punctuation()
    self.types  = self.remove_duplicated()
    self.token_ratio = self.get_token_ratio()

  def raw_full_text(self):
    raw_full_text = []
    for paragraph in self.document.paragraphs:
        raw_full_text.append(paragraph.text)
    return raw_full_text
  
  def trim_text(self, raw_text):
    trimmed_text = []

    for paragraph in raw_text:
      if paragraph != '':
          trimmed_text.append(paragraph)

    return trimmed_text

  
  def merge_string(self, raw_text):
    return ''.join(raw_text)

  
  def build_string(self):
    return self.merge_string(
      self.trim_text(
        self.raw_full_text()
      )
    )

  def build_tokens(self):
    return word_tokenize(self.string)

  def remove_punctuation(self):
    return [word.lower() for word in self.tokens if re.search("\w", word)]

  def remove_duplicated(self):
    return set(self.words)

  def token_count(self):
    return len(self.tokens)

  def word_count(self):
    return len(self.words)

  def type_count(self):
    return len(self.types)

  def get_token_ratio(self):
    return self.token_count() / self.type_count()

