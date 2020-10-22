import docx
from util.file_manager import file_extension
from util.file_manager import get_filename
from util.other_docs import OtherDoc
import util.log as log
import re

class Document():

  type_manager = None
  document = None
  string = None

  def __init__(self, path):
    self.document = self.__get_document(path)
    self.string = self.document.string

  def __get_document(self, path):
    if file_extension(path) == '.doc':
      return Doc(path)
    return Docx(path)


class Docx():

  string = None
  document = None

  def __init__(self, path):

    log.debug('Trying to read: {}'.format(get_filename(path)))
    self.document = self.read_document(path)
    self.string = self.build_string()

  def read_document(self, path):
    return docx.Document(path)

  def raw_full_text(self):
    raw_full_text = []
    for paragraph in self.document.paragraphs:
        raw_full_text.append(paragraph.text)
    return raw_full_text
  
  def fix_void_paragraph(self, raw_text):
    trimmed_text = []

    for paragraph in raw_text:
      if paragraph != '':
          trimmed_text.append(paragraph)

    return trimmed_text
  
  def merge_string(self, raw_text):
    return ''.join(raw_text)


  def separate_glued_words(self, string):
    tokens = string.split(' ')
    for i, token in enumerate(tokens):
        if not re.search('http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', token):
            tokens[i] = tokens[i].replace('.', '. ')
    return ' '.join(tokens)

  def build_string(self):
    return self.separate_glued_words(
        self.merge_string(
            self.fix_void_paragraph(
                self.raw_full_text()
            )
        )
    )
    


class Doc(OtherDoc):
  def __init__(self, path):
    log.debug('Trying to read: {}'.format(get_filename(path)))
    super().__init__(path, '.doc')


class Rtf(OtherDoc):
  def __init__(self, path):
    log.debug('Trying to read: {}'.format(get_filename(path)))
    super().__init__(path, '.rtf')



