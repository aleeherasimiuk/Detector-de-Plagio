import docx
from util.file_manager import file_extension
from util.file_manager import get_filename
from util.generic_document import GenericDocument
from util.data_cleaning import separate_glued_words
from util.data_cleaning import merge_string
import util.log as log
import re

class WordDocument():

  type_manager = None
  document = None
  string = None
  paragraphs = []

  def __init__(self, path):
    self.document = self.__get_document(path)
    self.string = self.document.string
    self.paragraphs = self.document.paragraphs

  def __get_document(self, path):
    if file_extension(path) == '.doc':
      return Doc(path)
    return Docx(path)


class Docx():

  string = None
  document = None
  raw_text = None
  paragraphs = []

  def __init__(self, path):

    log.debug('Trying to read: {}'.format(get_filename(path)))
    self.document = self.read_document(path)
    self.raw_text = self.raw_full_text()
    self.string = self.build_string(self.raw_text)
    self.paragraphs = self.get_paragraphs()

  def read_document(self, path):
    return docx.Document(path)

  def raw_full_text(self):
    return [paragraph.text for paragraph in self.document.paragraphs]
  
  def fix_void_paragraph(self, raw_text):
    return [paragraph for paragraph in raw_text if paragraph != '']
  
  def build_string(self, raw_text):
    fixed_void_paragraph = self.fix_void_paragraph(raw_text)
    merged_string = merge_string(fixed_void_paragraph)
    separated_glued_words = separate_glued_words(merged_string)

    return separated_glued_words

  def get_paragraphs(self):
    return self.raw_text.copy()



class Doc(GenericDocument):
  def __init__(self, path):
    log.debug('Trying to read: {}'.format(get_filename(path)))
    super().__init__(path, '.doc')

  def get_paragraphs(self):
    return re.split('\.\n\n', self.string)


class Rtf(GenericDocument):
  def __init__(self, path):
    log.debug('Trying to read: {}'.format(get_filename(path)))
    super().__init__(path, '.rtf')

  def get_paragraphs(self):
    return re.split('\.\n', self.string)



